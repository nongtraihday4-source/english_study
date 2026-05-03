#!/usr/bin/env python3
"""
scripts/extract_grammar_docx.py
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Trích xuất nội dung Ngữ pháp từ 5 file DOCX (A1-C1) và
xuất ra các file JSON sẵn sàng cho lệnh seed_grammar.

Cách chạy (từ thư mục gốc dự án):
    pip install python-docx
    python scripts/extract_grammar_docx.py

Output:
    scripts/grammar_fixtures/grammar_A1.json
    scripts/grammar_fixtures/grammar_A2.json
    scripts/grammar_fixtures/grammar_B1.json
    scripts/grammar_fixtures/grammar_B2.json
    scripts/grammar_fixtures/grammar_C1.json
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""
import json
import os
import re
import sys

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx không được cài. Chạy: pip install python-docx")
    sys.exit(1)

# ─── Paths ────────────────────────────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SOURCE_DIR = os.path.join(SCRIPT_DIR, "..", "old_system", "source")
OUTPUT_DIR = os.path.join(SCRIPT_DIR, "grammar_fixtures")

DOCX_FILES = {
    "A1": "Complete list of A1 grammar contents.docx",
    "A2": "Complete list of A2 grammar contents.docx",
    "B1": "Complete list of B1 grammar contents.docx",
    "B2": "Complete list of B2 grammar contents.docx",
    "C1": "Complete list of C1 grammar contents.docx",
}

# ─── Helpers ──────────────────────────────────────────────────────────────────

def _slugify(text: str) -> str:
    """Simple ASCII slug."""
    text = text.lower().strip()
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"[\s_-]+", "-", text)
    text = re.sub(r"^-+|-+$", "", text)
    return text[:120]


def _detect_formula(text: str) -> str:
    """
    Heuristically detect grammar formula patterns in a line.
    E.g.: 'S + V(s/es) + O', 'Subject + am/is/are + V-ing'
    """
    formula_patterns = [
        r"S\s*\+",            # S + V
        r"Subject\s*\+",
        r"\b(am|is|are|was|were)\b.*\+",
        r"V\s*[-/]ing",       # V-ing
        r"\bwh[-\s]",         # wh-question
        r"\bdo\s*/\s*does\b",
        r"\bdid\b.*\?",
        r"Modal\s*\+",
        r"have\s*/\s*has",
        r"to\s+be\s*\+",
    ]
    for p in formula_patterns:
        if re.search(p, text, re.IGNORECASE):
            return text.strip()
    return ""


def _split_sentence_translation(text: str):
    """
    Try to split an example line of form:
      'She is reading a book. → Cô ấy đang đọc sách.'
    or
      'She is reading a book. (Cô ấy đang đọc sách.)'
    Returns (sentence, translation) or (text, "").
    """
    # Arrow separator
    for sep in [" → ", " - ", " – ", " — "]:
        if sep in text:
            parts = text.split(sep, 1)
            return parts[0].strip(), parts[1].strip()
    # Vietnamese in parentheses
    m = re.search(r"\(([^)]+)\)\s*$", text)
    if m:
        translation = m.group(1).strip()
        sentence = text[: m.start()].strip()
        if re.search(r"[àáảãạăắặẳẵấầẩẫậđèéẻẽẹêếềểễệìíỉĩịòóỏõọôốồổỗộùúủũụưứừửữựỳýỷỹỵ]", translation):
            return sentence, translation
    return text.strip(), ""


# ─── Main Parser ──────────────────────────────────────────────────────────────

def _detect_doc_mode(doc) -> str:
    """
    Detect heading hierarchy used in the DOCX.

    'h1_mode'  (C1): Heading 1 = GrammarTopic, Heading 2 = GrammarRule
    'h2_mode'  (A1-B2): Heading 2 = GrammarTopic, Heading 4 = GrammarRule
                        Heading 3 = TOC categories (skip), Heading 6 = image captions (skip)
    """
    h1_topics = [
        p for p in doc.paragraphs
        if p.style and p.style.name == "Heading 1"
        and p.text.strip()
        and "complete list" not in p.text.lower()
        and "download" not in p.text.lower()
        and len(p.text.strip()) < 200
    ]
    return "h1_mode" if len(h1_topics) >= 3 else "h2_mode"


def _make_topic(text: str, level: str, order: int, chapter: str = "") -> dict:
    clean = re.sub(r"^\d+[\.\)]\s*", "", text).strip()
    return {
        "level": level,
        "title": clean,
        "slug": f"{level.lower()}-{_slugify(clean)}",
        "order": order,
        "chapter": chapter,
        "description": "",
        "analogy": "",
        "real_world_use": "",
        "memory_hook": "",
        "icon": _pick_icon(clean),
        "rules": [],
    }


def _make_rule(text: str, order: int) -> dict:
    clean = re.sub(r"^\d+[\.\)]\s*", "", text).strip()
    return {
        "title": clean,
        "formula": "",
        "explanation": "",
        "memory_hook": "",
        "is_exception": "exception" in text.lower() or "ngoại lệ" in text.lower(),
        "order": order,
        "examples": [],
    }


def _append_content(current_rule, current_topic, text: str):
    """Route body text to example list or explanation string."""
    if current_rule is None and current_topic is not None:
        if current_topic["description"]:
            current_topic["description"] += " " + text
        else:
            current_topic["description"] = text
        return

    if current_rule is not None:
        formula_text = _detect_formula(text)
        if formula_text and not current_rule["formula"]:
            current_rule["formula"] = formula_text
            return
        if _looks_like_example(text):
            sentence, translation = _split_sentence_translation(text)
            current_rule["examples"].append({
                "sentence": sentence,
                "translation": translation,
                "context": "",
                "highlight": _detect_highlight(sentence, current_rule["title"]),
            })
        else:
            if current_rule["explanation"]:
                current_rule["explanation"] += " " + text
            else:
                current_rule["explanation"] = text


def parse_docx(filepath: str, level: str) -> dict:
    """
    Parse a DOCX file and return a dict with 'chapters' and 'topics'.

    Two modes detected automatically:
      h1_mode (C1):   Heading 1 → topic, Heading 2 → rule, Heading 3 → chapter
      h2_mode (A1-B2): Heading 3 (TOC) → chapter list, Heading 2 → topic, Heading 4 → rule

    Chapter detection (h2_mode — two-pass):
      Pass 1: Collect canonical chapter names from the TOC section
              (H3 headings between 'Complete list…' H2 and the first real H2).
      Pass 2: Main parse. Only treat content-section H3 as a chapter boundary
              when the text is in the TOC chapter set. Subsection H3 within
              topics (e.g. 'Form', 'Use', "Let's") are silently ignored.
              The very first content group inherits toc_chapters_ordered[0].
    """
    doc = Document(filepath)
    mode = _detect_doc_mode(doc)

    # ── Pass 1 (h2_mode): collect canonical TOC chapter names ────────────────
    toc_chapters_ordered: list = []
    toc_chapters_set: set = set()
    if mode == "h2_mode":
        _in_toc = False
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            sname = (para.style.name if para.style else "") or ""
            if sname == "Heading 2":
                if "complete list" in text.lower():
                    _in_toc = True
                    continue
                break  # first real content H2 → end of TOC
            if sname == "Heading 3" and _in_toc and text not in toc_chapters_set:
                toc_chapters_ordered.append(text)
                toc_chapters_set.add(text)

    # ── Pass 2: main parse ────────────────────────────────────────────────────
    topics: list = []
    chapters_seen: dict = {}   # name → {name, slug, order, icon, description}
    chapter_order = 0
    current_chapter = ""
    current_topic = None
    current_rule = None
    topic_order = 0
    rule_order = 0
    in_toc = (mode == "h2_mode")
    first_content_h2_seen = False

    def _register_chapter(name: str):
        nonlocal current_chapter, chapter_order
        name = name.strip()
        if not name:
            return
        current_chapter = name
        if name not in chapters_seen:
            chapter_order += 1
            chapters_seen[name] = {
                "name": name,
                "slug": _slugify(name),
                "order": chapter_order,
                "icon": _pick_icon(name),
                "description": "",
            }

    def flush_rule():
        nonlocal current_rule
        if current_rule is not None and current_topic is not None:
            current_topic["rules"].append(current_rule)
        current_rule = None

    def flush_topic():
        nonlocal current_topic, current_rule, rule_order
        flush_rule()
        if current_topic is not None:
            topics.append(current_topic)
        current_topic = None
        rule_order = 0

    # ─────────────────────────────────────────────────────────────────────────
    if mode == "h1_mode":
        # C1: H1=topic, H2/H3/H4=rule (C1 DOCX has no TOC chapter structure)
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            sname = (para.style.name if para.style else "") or ""

            if sname == "Heading 6" or "download" in text.lower():
                continue

            if sname == "Heading 1":
                flush_topic()
                topic_order += 1
                rule_order = 0
                current_topic = _make_topic(text, level, topic_order, current_chapter)
                continue

            if sname in ("Heading 2", "Heading 3", "Heading 4"):
                if current_topic is None:
                    topic_order += 1
                    rule_order = 0
                    current_topic = _make_topic(f"{level} Grammar", level, topic_order, current_chapter)
                flush_rule()
                rule_order += 1
                current_rule = _make_rule(text, rule_order)
                continue

            if current_topic is None:
                continue
            _append_content(current_rule, current_topic, text)

    else:
        # A1-B2: H2=topic, H4=rule, H6=skip; H3 only from TOC list = chapter boundary
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            sname = (para.style.name if para.style else "") or ""

            if sname == "Heading 6" or "download" in text.lower():
                continue

            if sname == "Heading 2":
                if "complete list" in text.lower():
                    in_toc = True
                    continue
                # First real content H2: bootstrap chapter to first TOC entry
                if not first_content_h2_seen:
                    first_content_h2_seen = True
                    if not current_chapter and toc_chapters_ordered:
                        _register_chapter(toc_chapters_ordered[0])
                in_toc = False
                flush_topic()
                topic_order += 1
                rule_order = 0
                current_topic = _make_topic(text, level, topic_order, current_chapter)
                continue

            if sname == "Heading 3":
                # Only valid chapter boundary when text is in TOC chapter set
                if not in_toc and text in toc_chapters_set:
                    _register_chapter(text)
                # Otherwise (in_toc=True, or subsection H3) → ignore
                continue

            if sname == "Heading 4":
                if in_toc or current_topic is None:
                    continue
                flush_rule()
                rule_order += 1
                current_rule = _make_rule(text, rule_order)
                continue

            if in_toc or current_topic is None:
                continue
            _append_content(current_rule, current_topic, text)

    flush_topic()

    # Final chapter list: preserve TOC order for h2_mode
    if mode == "h2_mode" and toc_chapters_ordered:
        final_chapters = [
            chapters_seen[name]
            for name in toc_chapters_ordered
            if name in chapters_seen
        ]
    else:
        final_chapters = list(chapters_seen.values())

    return {
        "chapters": final_chapters,
        "topics": topics,
    }


def _looks_like_example(text: str) -> bool:
    """Heuristic: line looks like a grammar example sentence."""
    # Reject overly long paragraphs — those are explanations, not examples
    if len(text) > 300:
        return False
    # Has translation separator
    if any(sep in text for sep in [" → ", " – ", " — "]):
        return True
    # Short sentence ending with punctuation, starting with English word
    if text.endswith((".", "!", "?")) and 3 <= len(text.split()) <= 40:
        words = text.split()
        if words and re.match(r"^[A-Za-z\"\'(]", words[0]):
            return True
    return False


def _detect_highlight(sentence: str, rule_title: str) -> str:
    """Try to identify the key grammar structure to highlight."""
    title_lower = rule_title.lower()
    patterns = {
        "present simple": r"\b(go|goes|do|does|is|am|are|have|has)\b",
        "present continuous": r"\b(am|is|are)\s+\w+ing\b",
        "past simple": r"\b(was|were|went|did|had|got|made)\b",
        "past continuous": r"\b(was|were)\s+\w+ing\b",
        "present perfect": r"\b(have|has)\s+(been|\w+ed|\w+en)\b",
        "future": r"\b(will|going to|shall)\b",
        "modal": r"\b(can|could|must|should|would|may|might|shall)\b",
        "passive": r"\b(am|is|are|was|were|been)\s+\w+ed\b",
        "conditional": r"\b(if|would|could|should)\b",
    }
    for key, pattern in patterns.items():
        if key in title_lower:
            m = re.search(pattern, sentence, re.IGNORECASE)
            if m:
                return m.group(0)
    return ""


def _pick_icon(title: str) -> str:
    """Pick a thematic emoji based on title keywords."""
    title_lower = title.lower()
    icons = {
        "present": "⏰", "past": "🕐", "future": "🔮",
        "continuous": "🔄", "perfect": "✅", "passive": "🔃",
        "conditional": "🤔", "modal": "💭", "question": "❓",
        "negative": "❌", "noun": "📦", "verb": "⚡",
        "adjective": "🎨", "adverb": "💨", "article": "📌",
        "pronoun": "👤", "preposition": "📍", "conjunction": "🔗",
        "comparison": "⚖️", "reported": "💬", "relative": "🔗",
        "infinitive": "📐", "gerund": "🔁", "clause": "📝",
        "sentence": "✍️", "word order": "🔀",
    }
    for keyword, icon in icons.items():
        if keyword in title_lower:
            return icon
    return "📚"


# ─── Post-process: fill empty descriptions ───────────────────────────────────

def post_process(topics: list) -> list:
    """Fill empty descriptions from first rule explanation."""
    for topic in topics:
        if not topic["description"] and topic["rules"]:
            first_exp = topic["rules"][0].get("explanation", "")
            if first_exp:
                topic["description"] = first_exp[:200]
    return topics


# ─── Report ───────────────────────────────────────────────────────────────────

def print_summary(level: str, data: dict):
    topics = data.get("topics", data) if isinstance(data, dict) else data
    chapters = data.get("chapters", []) if isinstance(data, dict) else []
    total_rules = sum(len(t["rules"]) for t in topics)
    total_examples = sum(
        sum(len(r["examples"]) for r in t["rules"]) for t in topics
    )
    print(f"  [{level}] Chapters: {len(chapters):2d}  |  Topics: {len(topics):3d}  |  Rules: {total_rules:4d}  |  Examples: {total_examples:4d}")


# ─── Entry point ─────────────────────────────────────────────────────────────

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    print()
    print("━" * 60)
    print("  ENGLISH STUDY — Grammar DOCX Extractor")
    print("━" * 60)

    all_stats = []
    for level, filename in DOCX_FILES.items():
        filepath = os.path.join(SOURCE_DIR, filename)
        if not os.path.exists(filepath):
            print(f"  ⚠  [{level}] File not found: {filepath}")
            continue

        print(f"  ℹ  [{level}] Parsing: {filename}")
        data = parse_docx(filepath, level)
        data["topics"] = post_process(data["topics"])
        print_summary(level, data)

        output_path = os.path.join(OUTPUT_DIR, f"grammar_{level}.json")
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"  ✔  [{level}] Saved → {output_path}")
        print()

        all_stats.append((level, len(data["topics"]), len(data["chapters"])))

    print("━" * 60)
    print("  EXTRACTION COMPLETE")
    print("  Kiểm tra thủ công các file JSON trước khi seed:")
    for level, topic_count, chapter_count in all_stats:
        print(f"    scripts/grammar_fixtures/grammar_{level}.json  ({chapter_count} chapters, {topic_count} topics)")
    print()
    print("  Bước tiếp theo:")
    print("    cd backend")
    print("    pip install python-docx  # nếu chưa cài")
    print("    python manage.py makemigrations grammar")
    print("    python manage.py migrate")
    print("    python manage.py seed_grammar")
    print("━" * 60)


if __name__ == "__main__":
    main()
